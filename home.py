from dotenv import load_dotenv
import os
import streamlit as st
import mysql.connector
import requests
import toml
from openai import OpenAI
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
import time
from functools import lru_cache
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO
from streamlit_lottie import st_lottie
from streamlit_lottie import st_lottie_spinner
import json
import requests


load_dotenv()
API_KEY = os.environ['OPENAI_API_KEY']
llm = OpenAI(openai_api_key=API_KEY)


prompt_template = PromptTemplate(
    input_variables=['ingredients'],
    template="""Give me a short but complete filipino recipe that could be made using the following ingredients: {ingredients}. 
    Add 1-4 ingredients only and remove unedible or ingredients not for human consumptions
    format the ingredients output to a dotted list
    if name of a dish/ready made food is entered on the ingredients you must use it as is and create different recipe with it
    provide: Recipe Name:
            \n Cook time, 
            \n serving size: 1 person, 
            \n Ingredients,
            \n Instructions.
            \n enjoy!
    """
)

second_template = """Generate a short and complete filipino recipe based from the {meals} but make it a bit different.
                    It must be made of the ingredients from the {ingredients} but add 1-3 more ingredients.
                    if name of a dish is entered on the ingredients you must use it as is
                    remove the unedible or not for human consumption food
                    provide: Recipe Name:
                            \n Cook time, 
                            \n serving size: 1 person, 
                            \n Ingredients,
                            \n Instructions,
                            \n enjoy!
                """
second_template = PromptTemplate(
    input_variables=['meals','ingredients'],
    template=second_template
)

third_template = """Generate another short, easy, and complete filipino recipe different from the {second_meals} and {meals}.
                    It must be made of the ingredients from the {ingredients} but you can add 3-6 more ingredients.
                    remove the unedible or not for human consumption food
                    provide: Recipe Name:
                            \n Cook time, 
                            \n serving size: 1 person, 
                            \n Ingredients,
                            \n Instructions.
                            \n enjoy!
                """
third_template = PromptTemplate(
    input_variables=['second_meals','ingredients', 'meals'],
    template=third_template
)

meal_chain = LLMChain(
    llm=llm, 
    prompt=prompt_template,
    output_key="meals",
    verbose=True
)

second_chain= LLMChain(
    llm=llm,
    prompt=second_template,
    output_key="second_meals",
    verbose=True
)

third_chain= LLMChain(
    llm=llm,
    prompt=third_template,
    output_key="third_meals",
    verbose=True
)

overall_chain = SequentialChain(
    chains=[meal_chain, second_chain, third_chain],
    input_variables=['ingredients'],
    output_variables=["meals","second_meals","third_meals"],
    verbose=True
)


def connect_to_database():
    secrets = toml.load("streamlit/secrets.toml")
    mysql_config = secrets["connections"]["mysql"]
        
    return mysql.connector.connect(
        host=mysql_config["host"],
        port=mysql_config["port"],
        user=mysql_config["user"],
        password=mysql_config["password"],
        database=mysql_config["database"]
    )

@st.cache_data
def load_lottiefile(filepath: str):
    with open(filepath, "r") as f:
        data = json.load(f)
    return data


def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()


def get_user_info():
    user_id = st.session_state.user_id
    email = st.session_state.email
    return user_id, email
    

def get_username():
    username = st.session_state.username
    return username

def generate_recipes(user_prompt):
    output = overall_chain({'ingredients': user_prompt})
    return output



def add_to_favourite(user_prompt, meal_recipe_name, output_meals):
    user_id, email = get_user_info()

    conn = connect_to_database()
    try:
        cursor = conn.cursor()

        cursor.execute("SHOW TABLES LIKE 'user_favorite_recipes'")
        table_exists = cursor.fetchone()
        if not table_exists:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_favorite_recipes (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    user_id INT,
                    user_prompt TEXT,
                    recipe_name TEXT,
                    recipe_info TEXT,
                    FOREIGN KEY (user_id) REFERENCES users(id)
                )
            """)                               
            conn.commit()

        # Check if the recipe already exists in user's favorites
        cursor.execute("SELECT COUNT(*) FROM user_favorite_recipes WHERE user_id = %s AND recipe_name = %s", (user_id, meal_recipe_name))
        existing_recipe_count = cursor.fetchone()[0]
        
        if existing_recipe_count > 0:
            st.warning('This recipe is already in your favorites.')
        else:
            query = "INSERT INTO user_favorite_recipes (user_id, user_prompt, recipe_name, recipe_info) VALUES (%s, %s, %s, %s)"
            cursor.execute(query, (user_id, user_prompt, meal_recipe_name, output_meals))
            conn.commit()
            st.success('Added to Favorites')
            time.sleep(1)

    except mysql.connector.Error as err:
        st.error(f"Error: {err}")
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn.is_connected(): 
            conn.close()
                    


def add_to_favourite_recommend(user_prompt, meal_recipe_name, output_meals):
    user_id, email = get_user_info()

    conn = connect_to_database()
    try:
        cursor = conn.cursor()

        # Check if the recipe already exists in user's favorites
        cursor.execute("SELECT COUNT(*) FROM user_favorite_recipes WHERE user_id = %s AND recipe_name = %s", (user_id, meal_recipe_name))
        existing_recipe_count = cursor.fetchone()[0]
        
        if existing_recipe_count > 0:
            st.warning('This recipe is already in your favorites.')
        else:
            query = "INSERT INTO user_favorite_recipes (user_id, user_prompt, recipe_name, recipe_info) VALUES (%s, %s, %s, %s)"
            cursor.execute(query, (user_id, user_prompt, meal_recipe_name, output_meals))
            conn.commit()
            st.success('Added to Favorites')
            time.sleep(1)  

    except mysql.connector.Error as err:
        st.error(f"Error: {err}")
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn.is_connected(): 
            conn.close()
                    
    



def rating_to_stars(rating):
    return "⭐" * rating

def get_comments_for_recipe(meal_recipe_name):
    conn = connect_to_database()
    
    cursor = conn.cursor()
    
    query = "SELECT email, rating FROM ratings WHERE recipe_name = (%s)"
    cursor.execute(query, (meal_recipe_name,))    
    comments_for_recipe = cursor.fetchall()
    conn.close()
    comments_texts = [(comment[0], rating_to_stars(comment[1])) for comment in comments_for_recipe]

    return comments_texts

def display_favorite_recipes(user_prompt):
    user_id, email = get_user_info()

    conn = connect_to_database()
    try:
        cursor = conn.cursor(dictionary=True)

        # Combine information from both tables
        cursor.execute("""
            SELECT ufr.recipe_name, ufr.recipe_info, r.rating
            FROM user_favorite_recipes ufr
            LEFT JOIN ratings r ON ufr.recipe_name = r.recipe_name
            WHERE ufr.user_prompt = %s
        """, (user_prompt,))
        
        result = cursor.fetchall()
        return result

    except mysql.connector.Error as err:
        st.error(f"Error: {err}")
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn.is_connected(): 
            conn.close()





# Function to get or create session state
def get_session_state():
    if 'session_state' not in st.session_state:
        st.session_state.session_state = {
            'generated_recipes': {},
            'generated_prompts': {},
            'last_updated': time.time()
        }
    else:
        # Check if session data is older than 10 seconds
        if time.time() - st.session_state.session_state['last_updated'] > 10:
            st.session_state.session_state = {
                'generated_recipes': {},
                'generated_prompts': {},
                'last_updated': time.time()
            }
    return st.session_state.session_state



# Function to add recipe to session state
def add_to_session_state(prompt, recipes):
    session_state = get_session_state()
    session_state['generated_prompts'][prompt] = time.time()
    session_state['generated_recipes'][prompt] = recipes



# Function to get recipe from session state
def get_recipe_from_session(prompt):
    session_state = get_session_state()
    if prompt in session_state['generated_prompts']:
        # Check if prompt is older than 10 seconds
        if time.time() - session_state['generated_prompts'][prompt] > 20:
            # Remove expired prompt and its recipe
            del session_state['generated_prompts'][prompt]
            del session_state['generated_recipes'][prompt]
            return None
        else:
            return session_state['generated_recipes'].get(prompt, None)
    else:
        return None



def generate_recipe_pdf(recipe_info, recipe_name):
    # Create a BytesIO buffer to store the PDF content
    from io import BytesIO
    buffer = BytesIO()
    
    # Create a canvas object with a larger width to accommodate longer text
    from reportlab.lib.pagesizes import letter, landscape
    c = canvas.Canvas(buffer, pagesize=landscape(letter))
    
    # Split recipe_info into lines
    lines = recipe_info.split("\n")
    
    # Calculate the height required for the text
    line_height = 15  # Adjust as needed
    total_height = len(lines) * line_height
    
    # Set the canvas size
    c.setPageSize((800, total_height))
    
    # Write recipe information to the PDF
    y = total_height - line_height  # Initial y coordinate
    for line in lines:
        c.drawString(100, y, line)
        y -= line_height  # Move to the next line
    
    # Save the PDF content
    c.save()
    
    # Get the value from the BytesIO buffer
    pdf_data = buffer.getvalue()
    
    # Close the buffer
    buffer.close()
    
    return pdf_data



def generate_word_document(recipe_name, recipe_info):
    doc = Document()
    doc.add_heading(recipe_name, 0)
    doc.add_paragraph(recipe_info)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def app():
    col1 , col2, col3 = st.columns(3)
    col1 , col2, col3 = st.columns(3)
    with col1:
        lottie_file2 = './Animations/animation-1.json'
        lottie_json2 = load_lottiefile(lottie_file2)
        st_lottie(
            lottie_json2,
            speed    = 1,
            reverse  = False,
            loop     = True,
            quality  = "low",
            height   = 250,
            width    = 250,
            key      = 'animation1', 
        )    
            
    with col2:
        lottie_file2 = './Animations/animation-2.json'
        lottie_json2 = load_lottiefile(lottie_file2)
        st_lottie(
            lottie_json2,
            speed    = 1,
            reverse  = False,
            loop     = True,
            quality  = "low",
            height   = 200,
            width    = 250,
            key      = 'animation2', 
        )

    with col3:
        lottie_file3 = './Animations/animation-3.json'
        lottie_json3 = load_lottiefile(lottie_file3)
        st_lottie(
            lottie_json3,
            speed    = 1,
            reverse  = False,
            loop     = True,
            quality  = "low",
            height   = 250,
            width    = 250,
            key      = 'animation3', 
        )
    st.subheader("From Waste to Taste:  AI-Enhanced Recipe for Leftovers")

    st.markdown("<hr>", unsafe_allow_html=True)

    user_id, email = get_user_info()

    session_state = st.session_state
    if 'recipe_prompt' not in session_state:
        session_state.recipe_prompt = None
    if 'output' not in session_state:
        session_state.output = None

    user_prompt = st.text_input("Enter a comma-separated list of ingredients")
    generate = st.button("Generate", key="Generate-Button")

    if generate and user_prompt:
        time.sleep(3)  # Delay for 3 seconds
        with st.spinner("Generating..."):
            output = generate_recipes(user_prompt)
            session_state.output = output
            session_state.recipe_prompt = user_prompt

    if session_state.output:
        output = session_state.output
        meal_recipe_name = output['meals'].split("Recipe Name:")[1].split("Cook time:")[0].strip()
        second_meal_recipe_name = output['second_meals'].split("Recipe Name:")[1].split("Cook time:")[0].strip()
        third_meal_recipe_name = output['third_meals'].split("Recipe Name:")[1].split("Cook time:")[0].strip()

        with st.expander(label=meal_recipe_name):
            st.write(output['meals'])
            output_meals = output['meals']
            col1, col2, col3 = st.columns(3)
            with col2:
                pass

            with col1:
                key_values = "first_meal"
                save_button = st.button('Save to Favorites', key=key_values)
                if save_button:
                    add_to_favourite(user_prompt, meal_recipe_name, output_meals)

                doc_io = generate_word_document(meal_recipe_name, output_meals)
                download_button_1  = st.download_button(label='Download Recipe', data=doc_io, file_name=f"{meal_recipe_name}.docx")
                if download_button_1:
                    pass
                    

        with st.expander(label=second_meal_recipe_name):
            st.write(output['second_meals'])
            output_meals = output['second_meals']
            col1, col2 = st.columns(2)
            with col2:
                pass
            with col1:
                key_values = "second_meal"
                save_button = st.button('Save to Favorites', key=key_values)
                if save_button:
                    add_to_favourite(user_prompt, second_meal_recipe_name, output_meals)

                doc_io = generate_word_document(second_meal_recipe_name, output_meals)
                download_button_2 = st.download_button(label='Download Recipe', data=doc_io, file_name=f"{second_meal_recipe_name}.docx")
                if download_button_2:
                    pass
        with st.expander(label=third_meal_recipe_name):
            st.write(output['third_meals'])
            output_meals = output['third_meals']
            col1, col2 = st.columns(2)
            with col2:
                pass

            with col1:
                key_values = "third_meal"
                save_button = st.button('Save to Favorites', key=key_values)
                if save_button:
                    add_to_favourite(user_prompt, third_meal_recipe_name, output_meals)
                
                doc_io = generate_word_document(third_meal_recipe_name, output_meals)
                download_button_3 = st.download_button(label='Download Recipe', data=doc_io, file_name=f"{third_meal_recipe_name}.docx")
                if download_button_3:
                    pass                       
        result = display_favorite_recipes(user_prompt)
        st.session_state.recommended_recipes = result
        st.session_state.recommended_recipes_expires_at = time.time() + 60  # Store for 1 minute

        if 'recommended_recipes' in st.session_state and 'recommended_recipes_expires_at' in st.session_state:
            if time.time() < st.session_state.recommended_recipes_expires_at:
                result = st.session_state.recommended_recipes
                st.subheader("Recommended Recipes")
                for item in result:
                    recipe_name  = item['recipe_name']
                    recipe_info  = item['recipe_info']

                    # Check if the user has already rated this recipe
                    conn = connect_to_database()
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM ratings WHERE user_id = %s AND recipe_name = %s", (user_id, recipe_name))
                    existing_ratings_count = cursor.fetchone()[0]
                    cursor.close()
                    conn.close()

                    # Check if rating is not None and greater than 3 before displaying
                    if item['rating'] is not None and int(item['rating']) >= 3:
                        rating_ = int(item['rating'])
                        rating  = rating_to_stars(rating_)
                        with st.expander(label=f"{recipe_name} {rating}"):
                            st.write(recipe_info)
                            if existing_ratings_count > 0:
                                st.success("You have already rated this recipe.")
                            else:
                                user_prompt =  None
                                save_button = st.button('Save to Favorites', key=f"Save_Button_Recipe{item}")
                                if save_button:
                                    add_to_favourite_recommend(user_prompt, recipe_name, recipe_info)
                    else:
                        pass

            else:
                st.write("Recommended recipes expired.")
        else:
            st.write("No recommended recipes found.")
